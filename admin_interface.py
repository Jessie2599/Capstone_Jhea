import tkinter as tk
from tkinter import ttk
from docx import Document
from tkinter import messagebox
from tkinter import messagebox, ttk
from PIL import Image, ImageTk
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import qrcode
from tkinter import filedialog
from tkcalendar import DateEntry
import keyboard
import pymysql
from tkinter import Frame
import sqlite3
import os  # Added for file handling
import shutil
import matplotlib.pyplot as plt
import numpy as np


PIN = "1234"
my_tree = None
rows = []

def connection():
    conn=pymysql.connect(host="localhost", user="root", password="", database="libtraq_db")
    return conn

def verify_pin(event=None):
    pin = pin_entry.get()
    if pin == PIN:
        open_next_window()
    else:
        messagebox.showerror("Invalid PIN", "Your PIN is invalid.")

def open_next_window():
    window.destroy()
    second_window()

def second_window():
    global my_tree
    def refreshTable():
        for date in my_tree.get_children():
            my_tree.delete(date)

        for array in read():
            my_tree.insert(parent='', index='end', iid=array, text="", values=(array), tag='orow')

        my_tree.tag_configure('orow', background='', font=("Arial", 10))
        my_tree.place(x=10, y=150)

    second_window = tk.Tk()
    second_window.geometry("1366x768")
    second_window.resizable(height=False, width=False)
    second_window.title("LibTraQ: Library Tracker and Monitoring System using QR Code")

    title_label = tk.Label(second_window, text="Library Tracker and Monitoring System Using QR Code", bg="Silver", font=("Arial Rounded MT Bold", 30, "bold"), borderwidth=10, relief="ridge")
    title_label.place(x=0, y=0, relwidth=0.99, height=100)

    label = tk.Label(second_window, text="Library Attendance", font=("Arial Bold", 20))
    label.place(x=5,y=110)
    def read():
        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM library_attendance")
        results = cursor.fetchall()
        conn.commit()
        conn.close()
        return results

    my_tree = ttk.Treeview(second_window,  height=25)

    my_tree['columns'] = ("ID Number", "First Name", "Middle Name", "Last Name", "Course", "Purpose", "Date & Time")

    my_tree.column("#0", width=0, stretch=tk.NO)
    my_tree.column("ID Number", anchor="center", width=160)
    my_tree.column("First Name", anchor="center", width=160)
    my_tree.column("Middle Name", anchor="center", width=160)
    my_tree.column("Last Name", anchor="center", width=160)
    my_tree.column("Course", anchor="center", width=150)
    my_tree.column("Purpose", anchor="center", width=200)
    my_tree.column("Date & Time", anchor="center", width=190)

    my_tree.heading("ID Number", text="ID Number", anchor="center")
    my_tree.heading("First Name", text="First Name", anchor="center")
    my_tree.heading("Middle Name", text="Middle Name", anchor="center")
    my_tree.heading("Last Name", text="Last Name", anchor="center")
    my_tree.heading("Course", text="Course", anchor="center")
    my_tree.heading("Purpose", text="Purpose", anchor="center")
    my_tree.heading("Date & Time", text="Date & Time", anchor="center")

    refreshTable()

    icon_image = Image.open("images/add_record.png")
    original_width, original_height = icon_image.size
    aspect_ratio = original_width / original_height
    new_width = 120
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    button_add_record = tk.Button(second_window, image=icon_photo, height=new_height, width=new_width, command=open_add_record_window)
    button_add_record.image = icon_photo
    button_add_record.place(x=1210, y=147)

    icon_image = Image.open("images/list_of_student_icon.png")
    new_width, new_height = 120, 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    list_of_students = tk.Button(second_window, image=icon_photo, command=open_list_of_students_window, height=new_height, width=new_width)
    list_of_students.image = icon_photo
    list_of_students.place(x=1210, y=237)

    icon_image = Image.open("images/generate_report_icon.png")
    original_width, original_height = icon_image.size
    aspect_ratio = original_width / original_height
    new_width = 120
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    button_generate_report = tk.Button(second_window, image=icon_photo, command=open_generate_report_window, height=new_height, width=new_width)
    button_generate_report.image = icon_photo
    button_generate_report.place(x=1210, y=327)

    icon_image = Image.open("images/library_utilization_icon.png")
    original_width, original_height = icon_image.size
    aspect_ratio = original_width / original_height
    new_width = 120
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    library_utilization_window = tk.Button(second_window, image=icon_photo, command=open_library_utilization_window, height=new_height, width=new_width)
    library_utilization_window.image = icon_photo
    library_utilization_window.place(x=1210, y=417)

    icon_image = Image.open("images/change_pin_icon.png")
    original_width, original_height = icon_image.size
    aspect_ratio = original_width / original_height
    new_width = 120
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    change_pin_window = tk.Button(second_window, image=icon_photo, command=open_change_pin_window, height=new_height, width=new_width)
    change_pin_window.image = icon_photo
    change_pin_window.place(x=1210, y=507)

    icon_image = Image.open("images/about_us_icon.png")
    original_width, original_height = icon_image.size
    aspect_ratio = original_width / original_height
    new_width = 120
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    about_us_window = tk.Button(second_window, image=icon_photo, command=open_about_us_window, height=new_height, width=new_width)
    about_us_window.image = icon_photo
    about_us_window.place(x=1210, y=597)

    second_window.mainloop()

def open_add_record_window():

    def add_record_window_save_to_db():
        id_get = id_entry.get()
        fname_get = first_name_entry.get()
        mname_get = middle_name_entry.get()
        lname_get = last_name_entry.get()
        sex_get = sex_var.get()
        course_get = course_var.get()
        status_get = status_var.get()

        try:
            conn = connection()
            cursor = conn.cursor()
            cursor.execute("INSERT INTO student (id_no, first_name, middle_name, last_name, sex, course, status) VALUES (%s, %s, %s, %s, %s, %s, %s)", (id_get, fname_get, mname_get, lname_get, sex_get, course_get, status_get))
            conn.commit()
            conn.close()
            messagebox.showinfo("Success", "Save Successful!")

            id_entry.delete(0, tk.END)
            first_name_entry.delete(0, tk.END)
            middle_name_entry.delete(0, tk.END)
            last_name_entry.delete(0, tk.END)
            sex_var.set("None")
            course_var.set("None")
            status_var.set("None")

        except Exception as error:
            messagebox.showerror("Error!", str(error))
            print(error)

    def generate_qr_code():
        qr_data = id_entry.get()

        for widget in qr_photo_box.winfo_children():
            widget.destroy()

        if qr_data:
            qr = qrcode.QRCode(version=1, box_size=8, border=1)
            qr.add_data(qr_data)
            qr.make(fit=True)
            qr_image = qr.make_image(fill="black", back_color="white")
            qr_image_tk = ImageTk.PhotoImage(qr_image)
            qr_label = tk.Label(qr_photo_box, image=qr_image_tk)
            qr_label.image = qr_image_tk
            qr_label.pack()

    def go_back():
        add_record_window.destroy()
        second_window.deiconify()

    add_record_window = tk.Toplevel()
    add_record_window.title("Add Record")
    add_record_window.geometry("1366x768")
    add_record_window.resizable(height=False, width=False)

    app_title_label = tk.Label(add_record_window, text="Library Tracker and Monitoring System Using QR Code", font=("Arial Rounded MT Bold", 30, "bold"))
    app_title_label.place(x=0, y=0, relwidth=0.99, height=100)

    title_label = tk.Label(add_record_window, text="Add Student Record", bg="gray", font=("Arial", 24, "bold"), borderwidth=1, relief="solid")
    title_label.place(x=0, y=102, relwidth=0.99, height=80)

    back_icon = ImageTk.PhotoImage(Image.open("images/back_icon.png"))
    back_button = tk.Button(add_record_window, image=back_icon, command=go_back, bd=0)
    back_button.place(x=10, y=123)

    id_label = tk.Label(add_record_window, text="ID Number:", font=("Arial", 18))
    id_label.place(x=520, y=200)

    id_entry = tk.Entry(add_record_window, width=50, bg="lightgray", font=("Arial", 18))
    id_entry.place(x=670, y=200)

    first_name_label = tk.Label(add_record_window, text="First Name:", font=("Arial", 18))
    first_name_label.place(x=520, y=240)

    first_name_entry = tk.Entry(add_record_window, width=50, bg="lightgray", font=("Arial", 18))
    first_name_entry.place(x=670, y=240)

    middle_name_label = tk.Label(add_record_window, text="Middle Name:", font=("Arial", 18))
    middle_name_label.place(x=520, y=280)

    middle_name_entry = tk.Entry(add_record_window, width=50, bg="lightgray", font=("Arial", 18))
    middle_name_entry.place(x=670, y=280)

    last_name_label = tk.Label(add_record_window, text="Last Name:", font=("Arial", 18))
    last_name_label.place(x=520, y=320)

    last_name_entry = tk.Entry(add_record_window, width=50, bg="lightgray", font=("Arial", 18))
    last_name_entry.place(x=670, y=320)

    sex_label = tk.Label(add_record_window, text="Sex:", font=("Arial", 18))
    sex_label.place(x=520, y=360)

    sex_var = tk.StringVar(add_record_window)
    sex_var.set("Male")

    male_radio = tk.Radiobutton(add_record_window, text="Male", variable=sex_var, value="Male", font=("Arial", 18))
    female_radio = tk.Radiobutton(add_record_window, text="Female", variable=sex_var, value="Female", font=("Arial", 18))

    male_radio.place(x=670, y=360)
    female_radio.place(x=770, y=360)

    course_label = tk.Label(add_record_window, text="Course:", font=("Arial", 18))
    course_label.place(x=520, y=400)

    course_var = tk.StringVar(add_record_window)
    course_var.set("BSA")

    bsa_radio = tk.Radiobutton(add_record_window, text="BSA", variable=course_var, value="BSA", font=("Arial", 18))
    bsed_radio = tk.Radiobutton(add_record_window, text="BSED", variable=course_var, value="BSED", font=("Arial", 18))
    beed_radio = tk.Radiobutton(add_record_window, text="BEED", variable=course_var, value="BEED", font=("Arial", 18))
    bshm_radio = tk.Radiobutton(add_record_window, text="BSHM", variable=course_var, value="BSHM", font=("Arial", 18))
    bsoa_radio = tk.Radiobutton(add_record_window, text="BSOA", variable=course_var, value="BSOA", font=("Arial", 18))
    bsit_radio = tk.Radiobutton(add_record_window, text="BSIT", variable=course_var, value="BSIT", font=("Arial", 18))

    bsa_radio.place(x=670, y=400)
    bsed_radio.place(x=770, y=400)
    beed_radio.place(x=870, y=400)
    bshm_radio.place(x=970, y=400)
    bsoa_radio.place(x=1070, y=400)
    bsit_radio.place(x=1170, y=400)

    status_label = tk.Label(add_record_window, text="Status:", font=("Arial", 18))
    status_label.place(x=520, y=440)

    status_var = tk.StringVar(add_record_window)
    status_var.set("Active")

    active_radio = tk.Radiobutton(add_record_window, text="Active", variable=status_var, value="Active", font=("Arial", 18))
    inactive_radio = tk.Radiobutton(add_record_window, text="Inactive", variable=status_var, value="Inactive", font=("Arial", 18))

    active_radio.place(x=670, y=440)
    inactive_radio.place(x=770, y=440)

    photo_box = Frame(add_record_window, width=220, height=200, bg="white", highlightbackground="black", highlightthickness=2)
    photo_box.place(x=10, y=200)

    qr_photo_box = Frame(add_record_window, width=220, height=200, bg="white", highlightbackground="black", highlightthickness=2)
    qr_photo_box.place(x=260, y=200)

    def upload_photo():
        file_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg *.gif *.bmp")])

        if file_path:
            try:
                image = Image.open(file_path)
                image.thumbnail((200, 200))  # Resize the image to fit the photo_box
                photo = ImageTk.PhotoImage(image)
                photo_label = tk.Label(photo_box, image=photo)
                photo_label.image = photo  # Keep a reference to avoid garbage collection
                photo_label.pack()
            except Exception as error:
                messagebox.showerror("Error!", str(error))

    upload_button = tk.Button(add_record_window, text="Upload Photo", font=("Arial", 18), command=upload_photo)
    upload_button.place(x=38, y=406)

    qr_code_button = tk.Button(add_record_window, text="Generate QR", font=("Arial", 18), command=generate_qr_code)
    qr_code_button.place(x=290, y=406)

    save_button = tk.Button(add_record_window, text="Save Record", bg="gray", font=("Arial", 18), command=add_record_window_save_to_db)
    save_button.place(x=670, y=490)

    add_record_window_mainloop()

def open_list_of_students_window():
    global table
    def update_student(event):
        global table
        selected_item = table.focus()
        if not selected_item:
            return
        selected_student_values = table.item(selected_item, 'values')

        search_entry.delete(0, tk.END)
        search_entry.insert(0, selected_student_values[1])


        update_window = tk.Toplevel()
        update_window.title("Update Student")
        update_window.geometry("605x440")

        id_no_label = tk.Label(update_window, text="ID Number:",  font=("Arial", 12))
        id_no_label.place(x=260, y=19)
        id_no_entry = tk.Entry(update_window,  width=28, bg="lightgray", font=("Arial", 10))
        id_no_entry.insert(0, selected_student_values[0])
        id_no_entry.place(x=200, y=45)

        first_name_label = tk.Label(update_window, text="First Name:",  font=("Arial", 12))
        first_name_label.place(x=260, y=72)
        first_name_entry = tk.Entry(update_window,width=28, bg="lightgray", font=("Arial", 10))
        first_name_entry.insert(0, selected_student_values[1])
        first_name_entry.place(x=200, y=100)

        middle_name_label = tk.Label(update_window, text="Middle Name:", font=("Arial", 12))
        middle_name_label.place(x=260, y=125)
        middle_name_entry = tk.Entry(update_window,width=28,  bg="lightgray", font=("Arial", 10))
        middle_name_entry.insert(0, selected_student_values[2])
        middle_name_entry.place(x=200, y=155)

        last_name_label = tk.Label(update_window, text="Last Name:", font=("Arial", 12))
        last_name_label.place(x=260, y=178)
        last_name_entry = tk.Entry(update_window, width=28,  bg="lightgray", font=("Arial", 10))
        last_name_entry.insert(0, selected_student_values[3])
        last_name_entry.place(x=200, y=205)

        sex_label = tk.Label(update_window, text="Sex:",  font=("Arial", 12))
        sex_label.place(x=260, y=231)
        sex_entry = tk.Entry(update_window, width=28,  bg="lightgray", font=("Arial", 10))
        sex_entry.insert(0, selected_student_values[4])
        sex_entry.place(x=200, y=255)

        course_label = tk.Label(update_window, text="Course:",  font=("Arial", 12))
        course_label.place(x=260, y=284)
        course_entry = tk.Entry(update_window,  width=28,  bg="lightgray", font=("Arial", 10))
        course_entry.insert(0, selected_student_values[5])
        course_entry.place(x=200, y=305)

        status_label = tk.Label(update_window, text="Status:",  font=("Arial", 12))
        status_label.place(x=260, y=333)
        status_entry = tk.Entry(update_window,  width=28,  bg="lightgray", font=("Arial", 10))
        status_entry.insert(0, selected_student_values[6])
        status_entry.place(x=200, y=355)

        update_button = tk.Button(update_window, bg="red", text="Update",  font=("Arial", 12), command=lambda: perform_update(update_window, id_no_entry, first_name_entry, middle_name_entry, last_name_entry, sex_entry, course_entry, status_entry))
        update_button.place(x=250, y= 390)

        photo_box = Frame(update_window, width=150, height=130, bg="white", highlightbackground="black",highlightthickness=2)
        photo_box.place(x=30, y=20)

        qr_photo_box = Frame(update_window, width=150, height=130, bg="white", highlightbackground="black",highlightthickness=2)
        qr_photo_box.place(x=420, y=20)

        upload_button = tk.Button(update_window, text="Upload Photo", bg= "gray", font=("Arial", 12))
        upload_button.place(x=50, y=155)

        qr_code_button = tk.Button(update_window, text="Generate QR",bg="gray", font=("Arial", 12))
        qr_code_button.place(x=440, y=155)

    def perform_update(update_window, id_no_entry, first_name_entry, middle_name_entry, last_name_entry, sex_entry, course_entry, status_entry):
        updated_id_no = id_no_entry.get()
        updated_first_name = first_name_entry.get()
        updated_middle_name = middle_name_entry.get()
        updated_last_name = last_name_entry.get()
        updated_sex = sex_entry.get()
        updated_course = course_entry.get()
        updated_status = status_entry.get()

        conn = connection()
        cursor = conn.cursor()
        update_query = "UPDATE student SET first_name = %s, middle_name = %s, last_name = %s, sex = %s, course = %s, status = %s WHERE id_no = %s"
        cursor.execute(update_query, (updated_first_name, updated_middle_name, updated_last_name, updated_sex, updated_course, updated_status, updated_id_no))
        conn.commit()
        conn.close()

        # Close the update window
        update_window.destroy()

        # Refresh the table with updated data
        populate_treeview(table)

    def populate_treeview(table):
        # Delete existing items in the Treeview
        for row in table.get_children():
            table.delete(row)

        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM student")
        rows = cursor.fetchall()
        conn.commit()
        conn.close()

        for row in rows:
            table.insert("", "end", values=row)

    def search_data():
        search_query = search_entry.get().strip().lower()

        for row in table.get_children():
            table.delete(row)

        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM student")
        rows = cursor.fetchall()

        for row in rows:
            if any(search_query in str(cell).strip().lower() for cell in row):
                table.insert("", "end", values=row)

        conn.close()

    def connection():
        conn = pymysql.connect(host="localhost", user="root", password="", database="libtraq_db")
        return conn

    def go_back():
        list_of_students_window.destroy()

    list_of_students_window = tk.Toplevel()
    list_of_students_window.title("List Of Students")
    list_of_students_window.geometry("1366x768")
    list_of_students_window.resizable(height=False, width=False)

    app_title_label = tk.Label(list_of_students_window, text="Library Tracker and Monitoring System Using QR Code", font=("Arial Rounded MT Bold", 30, "bold"))
    app_title_label.place(x=0, y=0, relwidth=0.99, height=100)

    title_label = tk.Label(list_of_students_window, text="List of Student", bg="gray", font=("Arial", 24, "bold"), borderwidth=1, relief="solid")
    title_label.place(x=0, y=102, relwidth=0.99, height=80)

    back_icon = ImageTk.PhotoImage(Image.open("images/back_icon.png"))
    back_button = tk.Button(list_of_students_window, image=back_icon, command=go_back, bd=0)
    back_button.place(x=10, y=123)

    search_image = Image.open("images/search_icon.png")
    search_photo = ImageTk.PhotoImage(search_image)

    search_button = tk.Button(list_of_students_window, image=search_photo, command=search_data)
    search_button.place(x=1270, y=190)

    search_entry = tk.Entry(list_of_students_window, font=("Arial", 25), width=69, bg="Light Grey")
    search_entry.place(x=10, y=205)

    table_frame = ttk.Frame(list_of_students_window)
    table_frame.place(x=5, y=270)

    table_columns = ("ID Number", "First Name", "Middle Name", "Last Name", "Sex", "Course", "Status", "QR Code", "Photo")
    table = ttk.Treeview(table_frame, columns=table_columns, show="headings", height=20)
    table.place(x=1000, y=1000)

    scrollbar = ttk.Scrollbar(list_of_students_window, orient="vertical", command=table.yview)
    table.configure(yscrollcommand=scrollbar.set)
    table.pack(fill="both", expand=False)
    scrollbar.pack(side="right", fill="y")

    table.bind("<Double-1>", update_student)

    for col in table_columns:
        table.heading(col, text=col)
        table.column(col, width=100)

    # Define column headers
    table.heading("ID Number", text="ID Number")
    table.heading("First Name", text="First Name")
    table.heading("Middle Name", text="Middle Name")
    table.heading("Last Name", text="Last Name")
    table.heading("Sex", text="Sex")
    table.heading("Course", text="Course")
    table.heading("Status", text="Status")
    table.heading("QR Code", text="QR Code")
    table.heading("Photo", text="Photo")

    # Set column widths
    table.column("ID Number", width=148, anchor="center")
    table.column("First Name", width=149, anchor="center")
    table.column("Middle Name", width=149, anchor="center")
    table.column("Last Name", width=149, anchor="center")
    table.column("Sex", width=148, anchor="center")
    table.column("Course", width=148, anchor="center")
    table.column("Status", width=148, anchor="center")
    table.column("QR Code", width=149, anchor="center")
    table.column("Photo", width=149, anchor="center")

    # Populate the Treeview with data
    populate_treeview(table)

    list_of_students_window.mainloop()

def open_generate_report_window():
    generate_report_window = tk.Toplevel()
    generate_report_window.title("Generate Report")
    generate_report_window.geometry("1366x768")
    generate_report_window.resizable(height=False, width=False)

    def refreshTable():
        for date in my_tree.get_children():
            my_tree.delete(date)

        for array in read():
            my_tree.insert(parent='', index='end', iid=array, values=array, tag='orow')
            my_tree.place(x=10, y=300)

    def read():
        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM library_attendance")
        results = cursor.fetchall()
        conn.close()
        return results

    def search_data():
        search_query = search_entry.get().strip().lower()

        for row in my_tree.get_children():
            my_tree.delete(row)

        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM library_attendance")
        rows = cursor.fetchall()

        for row in rows:
            if any(search_query in str(cell).strip().lower() for cell in row):
                my_tree.insert("", "end", values=row)

        conn.close()

    def connection():
        conn = pymysql.connect(host="localhost", user="root", password="", database="libtraq_db")
        return conn

    my_tree = ttk.Treeview(generate_report_window, height=20)

    my_tree['columns'] = ("ID Number", "First Name", "Middle Name", "Last Name", "Course", "Purpose", "Date & Time")

    my_tree.column("#0", width=0, stretch=tk.NO)
    my_tree.column("ID Number", anchor="center", width=180)
    my_tree.column("First Name", anchor="center", width=180)
    my_tree.column("Middle Name", anchor="center", width=180)
    my_tree.column("Last Name", anchor="center", width=180)
    my_tree.column("Course", anchor="center", width=170)
    my_tree.column("Purpose", anchor="center", width=220)
    my_tree.column("Date & Time", anchor="center", width=210)

    my_tree.heading("ID Number", text="ID Number", anchor="center")
    my_tree.heading("First Name", text="First Name", anchor="center")
    my_tree.heading("Middle Name", text="Middle Name", anchor="center")
    my_tree.heading("Last Name", text="Last Name", anchor="center")
    my_tree.heading("Course", text="Course", anchor="center")
    my_tree.heading("Purpose", text="Purpose", anchor="center")
    my_tree.heading("Date & Time", text="Date & Time", anchor="center")

    my_tree.pack()

    refreshTable()
    def go_back():
        generate_report_window.destroy()
    def generate_report():
        selected_courses = course_var.get()

    def print_report():
        selected_courses = course_var.get()

        # Create a Word document
        filename = "report.docx"
        document = Document()

        # Add content to the Word document (customize this part as needed)
        document.add_heading("Library Attendance Report", 0)
        document.add_paragraph("Selected Semester: " + semester_var.get())
        document.add_paragraph("Selected Courses: " + course_var.get())
        document.add_paragraph("Generated by: Mariter Asur")

        # Create a table with headers
        table = document.add_table(rows=1, cols=7)
        table.allow_autofit = False  # Prevent table from auto-adjusting column widths

        table.style = 'Table Grid'  # Apply a table style

        # Add table headers
        table_headers = table.rows[0].cells
        table_headers[0].text = "ID Number"
        table_headers[1].text = "First Name"
        table_headers[2].text = "Middle Name"
        table_headers[3].text = "Last Name"
        table_headers[4].text = "Course"
        table_headers[5].text = "Purpose"
        table_headers[6].text = "Date & Time"

        # Get data from the Treeview and add it to the Word document table
        for item in my_tree.get_children():
            values = my_tree.item(item, 'values')
            data = list(map(str, values))
            table.add_row().cells[0].text = data[0]
            table.rows[-1].cells[1].text = data[1]
            table.rows[-1].cells[2].text = data[2]
            table.rows[-1].cells[3].text = data[3]
            table.rows[-1].cells[4].text = data[4]
            table.rows[-1].cells[5].text = data[5]
            table.rows[-1].cells[6].text = data[6]

        # Save the Word document
        document.save(filename)

        # Display a messagebox
        tk.messagebox.showinfo("Print Successful", "The report has been generated and saved successfully.")

        # Move the file to the "Downloads" folder (change the path as needed)
        downloads_folder = os.path.expanduser("~") + "/Downloads"
        new_filepath = os.path.join(downloads_folder, filename)
        shutil.move(filename, new_filepath)

        print(f"Report generated and saved as {new_filepath}")

    def fetch_and_display_records(dummy_arg=None):
        selected_semester = semester_var.get()
        selected_course = course_var.get()

        # Clear the Treeview widget
        for row in my_tree.get_children():
            my_tree.delete(row)

        # Fetch records from the database based on the selected semester and course
        conn = connection()
        cursor = conn.cursor()

        # Define SQL query based on selected semester and course
        if selected_semester == "1st Semester":
            date_range_query = "date_and_time >= 'August' AND date_and_time <= 'December'"
        elif selected_semester == "2nd Semester":
            date_range_query = "date_and_time >= 'January' AND date_and_time <= 'June'"
        elif selected_semester == "Whole School Year":
            date_range_query = "(date_and_time >= 'August' AND date_and_time <= 'December') OR (date_and_time >= 'January' AND date_and_time <= 'June')"
        else:
            date_range_query = ""

        if selected_course == "All Courses":
            course_query = ""
        else:
            course_query = f"course = '{selected_course}'"

        # Construct the full SQL query based on semester and course
        if date_range_query and course_query:
            query = f"SELECT * FROM library_attendance WHERE {date_range_query} AND {course_query}"
        elif date_range_query:
            query = f"SELECT * FROM library_attendance WHERE {date_range_query}"
        elif course_query:
            query = f"SELECT * FROM library_attendance WHERE {course_query}"
        else:
            query = "SELECT * FROM library_attendance"

        cursor.execute(query)
        records = cursor.fetchall()
        conn.close()

        # Display fetched records in the Treeview widget
        for record in records:
            my_tree.insert("", "end", values=record)

    def generate_bar_graph():
        course_counts = {}

        # Count the attendance for each course
        for item in my_tree.get_children():
            values = my_tree.item(item, 'values')
            course = values[4]  # Assuming course is in the 5th column
            if course in course_counts:
                course_counts[course] += 1
            else:
                course_counts[course] = 1

        courses = list(course_counts.keys())
        counts = list(course_counts.values())

        plt.figure(figsize=(10, 6))
        plt.bar(courses, counts, color='blue')
        plt.xlabel('Course')
        plt.ylabel('Attendance Count')
        plt.title('Attendance by Course')
        plt.xticks(rotation=45)
        plt.tight_layout()
        plt.show()

    app_title_label = tk.Label(generate_report_window, text="Library Tracker and Monitoring System Using QR Code", font=("Arial Rounded MT Bold", 30, "bold"))
    app_title_label.place(x=0, y=0, relwidth=0.99, height=100)

    title_label = tk.Label(generate_report_window, text="Generate Report", bg="gray", font=("Arial", 24, "bold"), borderwidth=1, relief="solid")
    title_label.place(x=0, y=102, relwidth=0.99, height=80)

    title_label = tk.Label(generate_report_window, text="Generate Report", bg="gray", font=("Arial", 24, "bold"), borderwidth=1, relief="solid")
    title_label.place(x=0, y=102, relwidth=0.99, height=80)

    semester_label = tk.Label(generate_report_window, text="Select Semester:", font=("Arial", 18))
    semester_label.place(x=0, y=200)

    semester_var = tk.StringVar(generate_report_window)
    semester_choices = ["1st Semester", "2nd Semester", "Whole School Year"]
    semester_dropdown = tk.OptionMenu(generate_report_window, semester_var, *semester_choices, command=fetch_and_display_records)
    semester_dropdown.config(font=("Arial", 16), width=20)
    semester_dropdown.place(x=180, y=200)

    course_label = tk.Label(generate_report_window, text="Select Courses:", font=("Arial", 18))
    course_label.place(x=470, y=200)

    course_var = tk.StringVar(generate_report_window)
    course_choices = ["BSA", "BSED", "BEED", "BSHM", "BSOA", "BSIT", "All Courses"]
    course_dropdown = tk.OptionMenu(generate_report_window, course_var, *course_choices, command=fetch_and_display_records)
    course_dropdown.config(font=("Arial", 16), width=20)
    course_dropdown.place(x=640, y=200)

    print_preview_button = tk.Button(generate_report_window, bg="green", text="Print Preview", font=("Arial", 16), command=lambda: fetch_and_display_records(None))
    print_preview_button.place(x=1020, y=200)

    graph_button = tk.Button(generate_report_window, bg="orange", text="Graph", font=("Arial", 16),command=generate_bar_graph)
    graph_button.place(x=935, y=200)

    print_button = tk.Button(generate_report_window, text="Print", bg="red" ,font=("Arial", 16),command=print_report)
    print_button.place(x=1170, y=200)

    back_icon = ImageTk.PhotoImage(Image.open("images/back_icon.png"))
    back_button = tk.Button(generate_report_window, image=back_icon, command=go_back, bd=0)
    back_button.place(x=10, y=123)

    search_image = Image.open("images/search_icon.png")
    search_photo = ImageTk.PhotoImage(search_image)

    search_button = tk.Button(generate_report_window, image=search_photo, command=search_data)
    search_button.place(x=1245, y=221)

    search_entry = tk.Entry(generate_report_window, font=("Arial", 25), width=68, bg="Light Grey")
    search_entry.place(x=10, y=250)

    generate_report_window.mainloop()

def open_library_utilization_window():
    library_utilization_window = tk.Toplevel()
    library_utilization_window.title("Library Utilization")
    library_utilization_window.geometry("1366x768")
    library_utilization_window.resizable(height=False, width=False)

    def go_back():
        library_utilization_window.destroy()

    def generate_bar_graph():
        course_counts = {}

        # Count the attendance for each course
        for item in my_tree.get_children():
            values = my_tree.item(item, 'values')
            course = values[4]  # Assuming course is in the 5th column
            if course in course_counts:
                course_counts[course] += 1
            else:
                course_counts[course] = 1

        courses = list(course_counts.keys())
        counts = list(course_counts.values())

        fig, ax = plt.subplots(figsize=(6, 4))
        ax.bar(courses, counts, color='blue')
        ax.set_xlabel('Course')
        ax.set_ylabel('Attendance Count')
        ax.set_title('Attendance by Course')
        ax.tick_params(axis='x', rotation=45)

        # Create a FigureCanvasTkAgg widget to embed the graph in a Tkinter Canvas
        canvas = FigureCanvasTkAgg(fig, master=library_utilization_window)
        canvas.get_tk_widget().place(x=0, y=280)  # Adjust the coordinates as needed
        canvas.draw()

    def generate_leaderboard():
        attendance_counts = {}

        # Count the attendance for each individual
        for item in my_tree.get_children():
            values = my_tree.item(item, 'values')
            id_number = values[0]  # Assuming ID number is in the 1st column
            first_name = values[1]  # Assuming first name is in the 2nd column
            middle_name = values[2]  # Assuming middle name is in the 3rd column
            last_name = values[3]  # Assuming last name is in the 4th column
            course = values[4]  # Assuming course is in the 5th column

            student_key = (first_name, middle_name, last_name, course)

            if student_key in attendance_counts:
                attendance_counts[student_key] += 1
            else:
                attendance_counts[student_key] = 1

        # Sort the attendance counts in descending order and select the top 10
        sorted_attendance = sorted(attendance_counts.items(), key=lambda x: x[1], reverse=True)[:10]

        leaderboard_label = tk.Label(library_utilization_window, text="Frequent Visitor", font=("Arial", 20, "bold"))
        leaderboard_label.place(x=800, y=270)

        leaderboard_frame = tk.Frame(library_utilization_window)
        leaderboard_frame.place(x=610, y=320)

        # Create a scrollbar for the leaderboard
        scrollbar = tk.Scrollbar(leaderboard_frame, orient="vertical")
        scrollbar.pack(side="right", fill="y")

        # Create a listbox to display the leaderboard
        leaderboard_listbox = tk.Listbox(leaderboard_frame, yscrollcommand=scrollbar.set, font=("Arial", 20),
                                         selectmode="none", width=47, height=10)  # Adjust width and height as needed
        leaderboard_listbox.pack(fill="both", expand=True)

        scrollbar.config(command=leaderboard_listbox.yview)

        # Populate the leaderboard listbox with the top 10 individuals and their attendance counts
        for index, (student_key, attendance_count) in enumerate(sorted_attendance, start=1):
            first_name, middle_name, last_name, course = student_key

            # Include numbering before each name
            leaderboard_listbox.insert("end",
                                       f"{index}. {first_name} {middle_name} {last_name} ({course}): {attendance_count} attendance(s)")

    graph_button = tk.Button(library_utilization_window, bg="orange", text="Graph", font=("Arial", 16),
                             command=generate_bar_graph)
    graph_button.place(x=25, y=200)

    leaderboard_button = tk.Button(library_utilization_window, bg="orange", text="Leaderboard", font=("Arial", 16), command=generate_leaderboard)
    leaderboard_button.place(x=825, y=200)

    app_title_label = tk.Label(library_utilization_window, text="Library Tracker and Monitoring System Using QR Code", font=("Arial Rounded MT Bold", 30, "bold"))
    app_title_label.place(x=0, y=0, relwidth=0.99, height=100)

    title_label = tk.Label(library_utilization_window, text="Library Utilization", bg="gray", font=("Arial", 24, "bold"), borderwidth=1, relief="solid")
    title_label.place(x=0, y=102, relwidth=0.99, height=80)

    back_icon = ImageTk.PhotoImage(Image.open("images/back_icon.png"))
    back_button = tk.Button(library_utilization_window, image=back_icon, command=go_back, bd=0)
    back_button.place(x=10, y=123)

    library_utilization_window.mainloop()
def open_change_pin_window():
    change_pin_window = tk.Toplevel()
    change_pin_window.title("Change Pin")
    change_pin_window.geometry("1366x768")
    change_pin_window.resizable(height=False, width=False)

    def go_back():
        change_pin_window.destroy()
        second_window.deiconify()

    def change_pin():
        old_pin = old_pin_entry.get()
        new_pin = new_pin_entry.get()
        if not simulated and validate_pins(old_pin, new_pin):

            global current_pin
            current_pin = new_pin
            messagebox.showinfo("Success", "PIN changed successfully!")
            go_back()

    def simulate_button_press(button):
        button.invoke()

    def validate_pins(old_pin, new_pin):
        return True

    def simulate_enter(event=None):
        global simulated
        simulated = True

        change_pin()

    app_title_label = tk.Label(change_pin_window, text="Library Tracker and Monitoring System Using QR Code", font=("Arial Rounded MT Bold", 30, "bold"))
    app_title_label.place(x=0, y=0, relwidth=0.99, height=100)

    title_label = tk.Label(change_pin_window, text="Change Pin", bg="gray", font=("Arial", 24, "bold"), borderwidth=1, relief="solid")
    title_label.place(x=0, y=102, relwidth=0.99, height=80)

    back_icon = ImageTk.PhotoImage(Image.open("images/back_icon.png"))
    back_button = tk.Button(change_pin_window, image=back_icon, command=go_back, bd=0)
    back_button.place(x=10, y=123)

    old_pin_label = tk.Label(change_pin_window, text="Old Pin", font=("Arial", 18))
    old_pin_label.place(x=300, y=240)

    old_pin_entry = tk.Entry(change_pin_window, width=50, bg="lightgray", font=("Arial", 18))
    old_pin_entry.place(x=470, y=240)

    new_pin_label = tk.Label(change_pin_window, text="New Pin", font=("Arial", 18))
    new_pin_label.place(x=300, y=280)

    new_pin_entry = tk.Entry(change_pin_window, width=50, bg="lightgray", font=("Arial", 18))
    new_pin_entry.place(x=470, y=280)

    first_name_label = tk.Label(change_pin_window, text="First Name", font=("Arial", 18))
    first_name_label.place(x=300, y=320)

    first_name_entry = tk.Entry(change_pin_window, width=50, bg="lightgray", font=("Arial", 18))
    first_name_entry.place(x=470, y=320)

    middle_name_label = tk.Label(change_pin_window, text="Middle Name", font=("Arial", 18))
    middle_name_label.place(x=300, y=360)

    middle_name_entry = tk.Entry(change_pin_window, width=50, bg="lightgray", font=("Arial", 18))
    middle_name_entry.place(x=470, y=360)

    last_name_label = tk.Label(change_pin_window, text="Last Name", font=("Arial", 18))
    last_name_label.place(x=300, y=400)

    last_name_entry = tk.Entry(change_pin_window, width=50, bg="lightgray", font=("Arial", 18))
    last_name_entry.place(x=470, y=400)

    #sex_label = tk.Label(change_pin_window, text="Sex:", font=("Arial", 18))
    #sex_label.place(x=520, y=360)

    #sex_var = tk.StringVar(change_pin_window)
    #sex_var.set("Male")

    #male_radio = tk.Radiobutton(change_pin_window, text="Male", variable=sex_var, value="Male", font=("Arial", 18))
    #female_radio = tk.Radiobutton(change_pin_window, text="Female", variable=sex_var, value="Female",font=("Arial", 18))

    #male_radio.place(x=670, y=360)
    #female_radio.place(x=770, y=360)

    change_pin_window.bind("<Return>", simulate_enter)
    change_pin_window.protocol("WM_DELETE_WINDOW", go_back)

    simulated = False

    change_pin_button = tk.Button(change_pin_window, text="Change Pin", command=change_pin, font=("Arial", 24))
    change_pin_button.place(relx=0.5, rely=0.75, anchor=tk.CENTER)

    current_pin = "1234"

    change_pin_window.mainloop()
def open_about_us_window():
    about_us_window = tk.Toplevel()
    about_us_window.title("About Us")
    about_us_window.geometry("1366x768")
    about_us_window.resizable(height=False, width=False)

    def go_back():
        about_us_window.destroy()

    background_image = Image.open("images/about_us_background.png")
    background_image = background_image.resize((1366, 768), Image.LANCZOS)
    background_photo = ImageTk.PhotoImage(background_image)

    background_label = tk.Label(about_us_window, image=background_photo)
    background_label.place(x=0, y=0, relwidth=1, relheight=1)
    about_us_window.background = background_photo

    back_button = tk.Button(about_us_window, text="Back", command=go_back)
    back_button.place(relx=0, rely=0, anchor=tk.NW)

    about_us_window.mainloop()

def on_entry_click(event):
    if pin_entry.get() == hint_text:
        pin_entry.delete(0, "end")
        pin_entry.config(fg='black')

def on_focus_out(event):
    if not pin_entry.get():
        pin_entry.insert(0, hint_text)
        pin_entry.config(fg='grey')


if __name__ == '__main__':
    window = tk.Tk()
    window.title("LibTraQ: Library Tracker and Monitoring System using QR Code")
    window.geometry("1366x768")
    window.resizable(height=False, width=False)

    hint_text = "Enter PIN"

    background_image = tk.PhotoImage(file="images/admin_background.png")

    background_label = tk.Label(window, image=background_image)
    background_label.place(x=0, y=0, relwidth=1, relheight=1)

    pin_label = tk.Label(window, text="Enter PIN:", font=("Arial", 26), justify='center')
    pin_label.pack(pady=(540, 10))

    pin_entry = tk.Entry(window, show="*", font=("Arial", 30), fg='grey', justify='center')
    pin_entry.insert(0, hint_text)
    pin_entry.bind("<FocusIn>", on_entry_click)
    pin_entry.bind("<FocusOut>", on_focus_out)
    pin_entry.pack(pady=10)
    pin_entry.bind("<Return>", verify_pin)
    pin_entry.focus()

    login_button = tk.Button(window, text="Login", font=("Arial", 24), command=verify_pin)
    login_button.pack(pady=0)

    window.mainloop()
