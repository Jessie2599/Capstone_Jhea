import tkinter as tk
from tkinter import ttk
from docx import Document
from tkinter import messagebox
from tkinter import messagebox, ttk
from PIL import Image, ImageTk
import io
from functools import partial
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import qrcode
from tkinter import filedialog
from tkcalendar import DateEntry
import keyboard
import pymysql
from tkinter import Frame
import sqlite3
import os
import shutil
import matplotlib.pyplot as plt
import numpy as np
import win32print
import win32ui



PIN = "1234"
my_tree = None
rows = []

global_qr_code_image = None
global_photo_image = None

file_path = None


def connection():
    conn=pymysql.connect(host="localhost", user="root", password="", database="libtraq_db")
    return conn

def verify_pin(event=None):
    pin = pin_entry.get()
    if pin == PIN:
        open_next_window()
    else:
        messagebox.showerror("Invalid PIN", "Your PIN is invalid.")

def open_next_window():
    window.destroy()
    second_window()

def second_window():
    global my_tree
    def refreshTable():
        for date in my_tree.get_children():
            my_tree.delete(date)

        for array in read():
            my_tree.insert(parent='', index='end', iid=array, text="", values=(array), tag='orow')

        my_tree.tag_configure('orow', background='', font=("Arial", 10))
        my_tree.place(x=10, y=150)

    second_window = tk.Tk()
    second_window.geometry("1366x768")
    second_window.resizable(height=False, width=False)
    second_window.title("LibTraQ: Library Tracker and Monitoring System using QR Code")

    title_label = tk.Label(second_window, text="Library Tracker and Monitoring System Using QR Code", bg="Silver", font=("Arial Rounded MT Bold", 30, "bold"), borderwidth=10, relief="ridge")
    title_label.place(x=0, y=0, relwidth=0.99, height=100)

    label = tk.Label(second_window, text="Library Attendance", font=("Arial Bold", 20))
    label.place(x=5,y=110)
    def read():
        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM library_attendance")
        results = cursor.fetchall()
        conn.commit()
        conn.close()
        return results

    my_tree = ttk.Treeview(second_window,  height=27)
    my_tree['columns'] = ("ID Number", "First Name", "Middle Name", "Last Name", "Course", "Purpose", "Date & Time")

    my_tree.column("#0", width=0, stretch=tk.NO)
    my_tree.column("ID Number", anchor="center", width=160)
    my_tree.column("First Name", anchor="center", width=160)
    my_tree.column("Middle Name", anchor="center", width=160)
    my_tree.column("Last Name", anchor="center", width=160)
    my_tree.column("Course", anchor="center", width=150)
    my_tree.column("Purpose", anchor="center", width=200)
    my_tree.column("Date & Time", anchor="center", width=190)

    my_tree.heading("ID Number", text="ID Number", anchor="center")
    my_tree.heading("First Name", text="First Name", anchor="center")
    my_tree.heading("Middle Name", text="Middle Name", anchor="center")
    my_tree.heading("Last Name", text="Last Name", anchor="center")
    my_tree.heading("Course", text="Course", anchor="center")
    my_tree.heading("Purpose", text="Purpose", anchor="center")
    my_tree.heading("Date & Time", text="Date & Time", anchor="center")

    refreshTable()

    def confirm_logout():
        result = messagebox.askquestion("Logout", "Are you sure you want to logout?")
        if result == "yes":
            # Perform logout actions here, if needed
            second_window.destroy()

    custom_font = ("Helvetica", 29, "bold")  # Modify this to your desired font family, size, and style

    # Create a label with your desired text, set the custom font, and position it
    label_text = "ADMIN "
    label = tk.Label(second_window, text=label_text, font=custom_font)
    label.place(x=1210, y=110)

    label_text = "TAB "
    label = tk.Label(second_window, text=label_text, font=custom_font)
    label.place(x=1235, y=150)


    icon_image = Image.open("images/add_user_icon.png")
    original_width, original_height = icon_image.size
    aspect_ratio = original_width / original_height
    new_width = 120
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    button_add_record = tk.Button(second_window, image=icon_photo, height=new_height, width=new_width, command=open_add_record_window)
    button_add_record.image = icon_photo
    button_add_record.place(x=1210, y=200)

    icon_image = Image.open("images/list_of_student_icon.png")
    new_width, new_height = 120, 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    list_of_students = tk.Button(second_window, image=icon_photo, command=open_list_of_students_window, height=new_height, width=new_width)
    list_of_students.image = icon_photo
    list_of_students.place(x=1210, y=290)

    icon_image = Image.open("images/generate_report_icon.png")
    original_width, original_height = icon_image.size
    aspect_ratio = original_width / original_height
    new_width = 120
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    button_generate_report = tk.Button(second_window, image=icon_photo, command=open_generate_report_window, height=new_height, width=new_width)
    button_generate_report.image = icon_photo
    button_generate_report.place(x=1210, y=380)

    icon_image = Image.open("images/library_utilization_icon.png")
    original_width, original_height = icon_image.size
    aspect_ratio = original_width / original_height
    new_width = 120
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    library_utilization_window = tk.Button(second_window, image=icon_photo, command=open_library_utilization_window, height=new_height, width=new_width)
    library_utilization_window.image = icon_photo
    library_utilization_window.place(x=1210, y=470)


    icon_image = Image.open("images/about_us_icon.png")
    original_width, original_height = icon_image.size
    aspect_ratio = original_width / original_height
    new_width = 120
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    about_us_window = tk.Button(second_window, image=icon_photo, command=open_about_us_window, height=new_height, width=new_width)
    about_us_window.image = icon_photo
    about_us_window.place(x=1210, y=560)

    icon_image = Image.open("images/logout_icon.png")
    new_width = 120
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    # Create the logout button
    logout_button = tk.Button(second_window, image=icon_photo, height=new_height, width=new_width,
                              command=confirm_logout)
    logout_button.image = icon_photo
    logout_button.place(x=1210, y=650)

    second_window.mainloop()

def open_add_record_window():

    def generate_qr_code():
        global global_qr_code_image
        qr_data = id_entry.get()

        if qr_data:
            qr = qrcode.QRCode(version=1, box_size=8, border=1)
            qr.add_data(qr_data)
            qr.make(fit=True)
            qr_image = qr.make_image(fill="black", back_color="white")
            global_qr_code_image = ImageTk.PhotoImage(qr_image)

            for widget in qr_photo_box.winfo_children():
                widget.destroy()

            qr_label = tk.Label(qr_photo_box, image=global_qr_code_image)
            qr_label.image = global_qr_code_image
            qr_label.pack()

    def open_file_dialog():
        global file_path  # Access the global variable
        file_path = filedialog.askopenfilename(parent=add_record_window,
                                               filetypes=[("Image files", "*.jpg *.jpeg *.png *.gif")])
        if file_path:
            image = Image.open(file_path)
            image = image.resize((200, 180), Image.LANCZOS)
            global global_photo_image
            global_photo_image = ImageTk.PhotoImage(image)

            for widget in photo_box.winfo_children():
                widget.destroy()

            label = tk.Label(photo_box, image=global_photo_image)
            label.image = global_photo_image
            label.pack()

    def display_input_values(id_get, fname_get, mname_get, lname_get, sex_get, course_get, status_get):
        input_display_window = tk.Toplevel(add_record_window)
        input_display_window.geometry("600x250")  # Set a fixed size for the window

        # Create labels for each input value with adjusted font size
        tk.Label(input_display_window, text="Student ID:", font=("Arial", 10, "bold")).grid(row=0, column=0, padx=5,
                                                                                            pady=5, sticky='w')
        tk.Label(input_display_window, text=id_get, font=("Arial", 10)).grid(row=0, column=1, padx=5, pady=5,
                                                                             sticky='w')

        tk.Label(input_display_window, text="First Name:", font=("Arial", 10, "bold")).grid(row=1, column=0, padx=5,
                                                                                            pady=5, sticky='w')
        tk.Label(input_display_window, text=fname_get, font=("Arial", 10)).grid(row=1, column=1, padx=5, pady=5,
                                                                                sticky='w')

        tk.Label(input_display_window, text="Middle Name:", font=("Arial", 10, "bold")).grid(row=2, column=0, padx=5,
                                                                                             pady=5, sticky='w')
        tk.Label(input_display_window, text=mname_get, font=("Arial", 10)).grid(row=2, column=1, padx=5, pady=5,
                                                                                sticky='w')

        tk.Label(input_display_window, text="Last Name:", font=("Arial", 10, "bold")).grid(row=3, column=0, padx=5,
                                                                                           pady=5, sticky='w')
        tk.Label(input_display_window, text=lname_get, font=("Arial", 10)).grid(row=3, column=1, padx=5, pady=5,
                                                                                sticky='w')

        tk.Label(input_display_window, text="Sex:", font=("Arial", 10, "bold")).grid(row=4, column=0, padx=5, pady=5,
                                                                                     sticky='w')
        tk.Label(input_display_window, text=sex_get, font=("Arial", 10)).grid(row=4, column=1, padx=5, pady=5,
                                                                              sticky='w')

        tk.Label(input_display_window, text="Course:", font=("Arial", 10, "bold")).grid(row=5, column=0, padx=5, pady=5,
                                                                                        sticky='w')
        tk.Label(input_display_window, text=course_get, font=("Arial", 10)).grid(row=5, column=1, padx=5, pady=5,
                                                                                 sticky='w')

        tk.Label(input_display_window, text="Status:", font=("Arial", 10, "bold")).grid(row=6, column=0, padx=5, pady=5,
                                                                                        sticky='w')
        tk.Label(input_display_window, text=status_get, font=("Arial", 10)).grid(row=6, column=1, padx=5, pady=5,
                                                                                 sticky='w')
        # Display the QR code
        if global_qr_code_image:
            tk.Label(input_display_window, image=global_qr_code_image).grid(row=0, column=2, rowspan=7, padx=10,
                                                                            pady=10, sticky='w')

        # Display the photo
        if global_photo_image:
            tk.Label(input_display_window, image=global_photo_image).grid(row=0, column=3, rowspan=7, padx=10, pady=10,
                                                                          sticky='w')

    def add_record_window_save_to_db():
        id_get = id_entry.get()
        fname_get = first_name_entry.get()
        mname_get = middle_name_entry.get()
        lname_get = last_name_entry.get()
        sex_get = sex_var.get()
        course_get = course_var.get()
        status_get = status_var.get()

        if not global_qr_code_image:
            messagebox.showerror("Error!", "Please generate a QR code.", parent=add_record_window)
            return  # Stop execution if no QR code is generated

        if not global_photo_image:
            messagebox.showerror("Error!", "Please upload a photo.", parent=add_record_window)
            return  # Stop execution if no photo is uploaded

        try:
            conn = connection()
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO student (id_no, first_name, middle_name, last_name, sex, course, status, qr_code, photo) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)",
                (id_get, fname_get, mname_get, lname_get, sex_get, course_get, status_get, global_qr_code_image,
                 global_photo_image))
            conn.commit()
            conn.close()

            display_input_values(id_get, fname_get, mname_get, lname_get, sex_get, course_get, status_get)

            # Clear input fields
            id_entry.delete(0, tk.END)
            first_name_entry.delete(0, tk.END)
            middle_name_entry.delete(0, tk.END)
            last_name_entry.delete(0, tk.END)
            sex_var.set("None")
            course_var.set("None")
            status_var.set("None")

            # Destroy the images in the photo_box and qr_photo_box
            for widget in photo_box.winfo_children():
                widget.destroy()

            for widget in qr_photo_box.winfo_children():
                widget.destroy()

        except Exception as error:
            message = f"Error: {str(error)}"
            messagebox.showerror("Error!", message, parent=add_record_window)
            print(error)

    qr_code_dir = "qr_code_bin"
    if not os.path.exists(qr_code_dir):
        os.makedirs(qr_code_dir)

    def print_qr_code():
        qr_data = id_entry.get()

        if qr_data:
            qr = qrcode.QRCode(version=1, box_size=8, border=1)
            qr.add_data(qr_data)
            qr.make(fit=True)
            qr_image = qr.make_image(fill="black", back_color="white")

            # Save the QR code to a file in the qr_code_dir directory
            qr_code_filename = f"{qr_code_dir}/{qr_data}.png"
            qr_image.save(qr_code_filename)

            try:
                # Open the QR code image with the default image viewer
                os.system(qr_code_filename)

                messagebox.showinfo("Success",
                                    "QR code image opened for printing. Please use the image viewer's print option.",
                                    parent=add_record_window)  # Set the parent window here
            except Exception as e:
                messagebox.showerror("Error", f"Failed to open the image for printing: {str(e)}",
                                     parent=add_record_window)
    def go_back():
        add_record_window.destroy()

    add_record_window = tk.Toplevel()
    add_record_window.title("Add Record")
    add_record_window.geometry("1366x768")
    add_record_window.resizable(height=False, width=False)

    app_title_label = tk.Label(add_record_window, text="Library Tracker and Monitoring System Using QR Code", font=("Arial Rounded MT Bold", 30, "bold"))
    app_title_label.place(x=0, y=0, relwidth=0.99, height=100)

    title_label = tk.Label(add_record_window, text="Add Student Record", bg="gray", font=("Arial", 24, "bold"), borderwidth=1, relief="solid")
    title_label.place(x=0, y=102, relwidth=0.99, height=80)

    back_icon = ImageTk.PhotoImage(Image.open("images/back_icon.png"))
    back_button = tk.Button(add_record_window, image=back_icon, command=go_back, bd=0)
    back_button.place(x=10, y=123)

    id_label = tk.Label(add_record_window, text="ID Number:", font=("Arial", 18))
    id_label.place(x=520, y=200)

    id_entry = tk.Entry(add_record_window, width=50, bg="lightgray", font=("Arial", 18))
    id_entry.place(x=670, y=200)

    first_name_label = tk.Label(add_record_window, text="First Name:", font=("Arial", 18))
    first_name_label.place(x=520, y=240)

    first_name_entry = tk.Entry(add_record_window, width=50, bg="lightgray", font=("Arial", 18))
    first_name_entry.place(x=670, y=240)

    middle_name_label = tk.Label(add_record_window, text="Middle Name:", font=("Arial", 18))
    middle_name_label.place(x=520, y=280)

    middle_name_entry = tk.Entry(add_record_window, width=50, bg="lightgray", font=("Arial", 18))
    middle_name_entry.place(x=670, y=280)

    last_name_label = tk.Label(add_record_window, text="Last Name:", font=("Arial", 18))
    last_name_label.place(x=520, y=320)

    last_name_entry = tk.Entry(add_record_window, width=50, bg="lightgray", font=("Arial", 18))
    last_name_entry.place(x=670, y=320)

    sex_label = tk.Label(add_record_window, text="Sex:", font=("Arial", 18))
    sex_label.place(x=520, y=360)

    sex_var = tk.StringVar(add_record_window)
    sex_var.set("Male")

    male_radio = tk.Radiobutton(add_record_window, text="Male", variable=sex_var, value="Male", font=("Arial", 18))
    female_radio = tk.Radiobutton(add_record_window, text="Female", variable=sex_var, value="Female", font=("Arial", 18))

    male_radio.place(x=670, y=360)
    female_radio.place(x=770, y=360)

    course_label = tk.Label(add_record_window, text="Course:", font=("Arial", 18))
    course_label.place(x=520, y=400)

    course_var = tk.StringVar(add_record_window)
    course_var.set("BSA")

    bsa_radio = tk.Radiobutton(add_record_window, text="BSA", variable=course_var, value="BSA", font=("Arial", 18))
    bsed_radio = tk.Radiobutton(add_record_window, text="BSED", variable=course_var, value="BSED", font=("Arial", 18))
    beed_radio = tk.Radiobutton(add_record_window, text="BEED", variable=course_var, value="BEED", font=("Arial", 18))
    bshm_radio = tk.Radiobutton(add_record_window, text="BSHM", variable=course_var, value="BSHM", font=("Arial", 18))
    bsoa_radio = tk.Radiobutton(add_record_window, text="BSOA", variable=course_var, value="BSOA", font=("Arial", 18))
    bsit_radio = tk.Radiobutton(add_record_window, text="BSIT", variable=course_var, value="BSIT", font=("Arial", 18))

    bsa_radio.place(x=670, y=400)
    bsed_radio.place(x=770, y=400)
    beed_radio.place(x=870, y=400)
    bshm_radio.place(x=970, y=400)
    bsoa_radio.place(x=1070, y=400)
    bsit_radio.place(x=1170, y=400)

    status_label = tk.Label(add_record_window, text="Status:", font=("Arial", 18))
    status_label.place(x=520, y=440)

    status_var = tk.StringVar(add_record_window)
    status_var.set("Active")

    active_radio = tk.Radiobutton(add_record_window, text="Active", variable=status_var, value="Active", font=("Arial", 18))
    inactive_radio = tk.Radiobutton(add_record_window, text="Inactive", variable=status_var, value="Inactive", font=("Arial", 18))

    active_radio.place(x=670, y=440)
    inactive_radio.place(x=770, y=440)

    photo_box = Frame(add_record_window, width=220, height=200, bg="white", highlightbackground="black", highlightthickness=2)
    photo_box.place(x=10, y=200)

    qr_photo_box = Frame(add_record_window, width=220, height=200, bg="white", highlightbackground="black", highlightthickness=2)
    qr_photo_box.place(x=260, y=200)

    upload_button = tk.Button(add_record_window, text="Upload Photo", font=("Arial", 18), command=open_file_dialog)
    upload_button.place(x=38, y=406)

    qr_code_button = tk.Button(add_record_window, text="Generate QR", font=("Arial", 18), command=generate_qr_code)
    qr_code_button.place(x=290, y=406)

    qr_code_save = tk.Button(add_record_window, text="Print", font=("Arial", 18), command=print_qr_code)
    qr_code_save.place(x=350, y=500)

    save_button = tk.Button(add_record_window, text="Save Record", bg="gray", font=("Arial", 18), command=add_record_window_save_to_db)
    save_button.place(x=670, y=490)

    add_record_window.mainloop()


def open_list_of_students_window():
    global table

    def update_student(event):
        global table
        selected_item = table.focus()
        if not selected_item:
            return
        selected_student_values = table.item(selected_item, 'values')

        search_entry.delete(0, tk.END)
        search_entry.insert(0, selected_student_values[1])
        photo_data = selected_student_values[7]  # Assuming that the photo is in index 7

        if photo_data:
            # Encode the string data to bytes
            photo_data_bytes = photo_data.encode()

            # Create a BytesIO stream from the bytes
            photo_stream = io.BytesIO(photo_data_bytes)

            # Display the QR code

        update_window = tk.Toplevel()
        update_window.title("Update Student")
        update_window.geometry("605x440")

        id_no_label = tk.Label(update_window, text="ID Number:", font=("Arial", 12))
        id_no_label.place(x=260, y=19)
        id_no_entry = tk.Entry(update_window, width=28, bg="lightgray", font=("Arial", 10))
        id_no_entry.insert(0, selected_student_values[0])
        id_no_entry.place(x=200, y=45)

        first_name_label = tk.Label(update_window, text="First Name:", font=("Arial", 12))
        first_name_label.place(x=260, y=72)
        first_name_entry = tk.Entry(update_window, width=28, bg="lightgray", font=("Arial", 10))
        first_name_entry.insert(0, selected_student_values[1])
        first_name_entry.place(x=200, y=100)

        middle_name_label = tk.Label(update_window, text="Middle Name:", font=("Arial", 12))
        middle_name_label.place(x=260, y=125)
        middle_name_entry = tk.Entry(update_window, width=28, bg="lightgray", font=("Arial", 10))
        middle_name_entry.insert(0, selected_student_values[2])
        middle_name_entry.place(x=200, y=155)

        last_name_label = tk.Label(update_window, text="Last Name:", font=("Arial", 12))
        last_name_label.place(x=260, y=178)
        last_name_entry = tk.Entry(update_window, width=28, bg="lightgray", font=("Arial", 10))
        last_name_entry.insert(0, selected_student_values[3])
        last_name_entry.place(x=200, y=205)

        sex_label = tk.Label(update_window, text="Sex:", font=("Arial", 12))
        sex_label.place(x=260, y=231)
        sex_entry = tk.Entry(update_window, width=28, bg="lightgray", font=("Arial", 10))
        sex_entry.insert(0, selected_student_values[4])
        sex_entry.place(x=200, y=255)

        course_label = tk.Label(update_window, text="Course:", font=("Arial", 12))
        course_label.place(x=260, y=284)
        course_entry = tk.Entry(update_window, width=28, bg="lightgray", font=("Arial", 10))
        course_entry.insert(0, selected_student_values[5])
        course_entry.place(x=200, y=305)

        status_label = tk.Label(update_window, text="Status:", font=("Arial", 12))
        status_label.place(x=260, y=333)
        status_entry = tk.Entry(update_window, width=28, bg="lightgray", font=("Arial", 10))
        status_entry.insert(0, selected_student_values[6])
        status_entry.place(x=200, y=355)

        update_button = tk.Button(update_window, bg="red", text="Update", font=("Arial", 12),
                                  command=lambda: perform_update(update_window, id_no_entry, first_name_entry,
                                                                 middle_name_entry, last_name_entry, sex_entry,
                                                                 course_entry, status_entry))
        update_button.place(x=250, y=390)

        photo_box = tk.Canvas(update_window, width=150, height=130, bg="white", highlightbackground="black",
                          highlightthickness=2)
        photo_box.place(x=30, y=20)

        qr_photo_box = tk.Canvas(update_window, width=150, height=130, bg="white", highlightbackground="black",
                             highlightthickness=2)
        qr_photo_box.place(x=420, y=20)

        upload_button = tk.Button(update_window, text="Upload Photo", bg="gray", font=("Arial", 12))
        upload_button.place(x=50, y=155)

        qr_code_button = tk.Button(update_window, text="Generate QR", bg="gray", font=("Arial", 12))
        qr_code_button.place(x=440, y=155)


        if photo_data:
            # Convert binary photo data to an ImageTk.PhotoImage
            photo = Image.open(io.BytesIO(photo_data[0]))
            photo = photo.resize((250, 250), Image.LANCZOS)
            photo = ImageTk.PhotoImage(photo)

            # Create a label to display the photo in the photo_box
            photo_label = tk.Label(photo_box, image=photo)
            photo_label.image = photo  # Keep a reference to prevent it from being garbage collected
            photo_label.pack(fill="both", expand=True)


        if qr_data:
            # Convert binary QR data to an ImageTk.PhotoImage
            qr_code = Image.open(io.BytesIO(qr_data[0]))
            qr_code = qr_code.resize((200, 200), Image.LANCZOS)
            qr_code = ImageTk.PhotoImage(qr_code)

            # Create a label to display the QR code in the qr_box
            qr_label = tk.Label(qr_box, image=qr_code)
            qr_label.image = qr_code  # Keep a reference to prevent it from being garbage collected
            qr_label.pack(fill="both", expand=True)

    def perform_update(update_window, id_no_entry, first_name_entry, middle_name_entry, last_name_entry, sex_entry,
                       course_entry, status_entry):
        updated_id_no = id_no_entry.get()
        updated_first_name = first_name_entry.get()
        updated_middle_name = middle_name_entry.get()
        updated_last_name = last_name_entry.get()
        updated_sex = sex_entry.get()
        updated_course = course_entry.get()
        updated_status = status_entry.get()

        conn = connection()
        cursor = conn.cursor()
        update_query = "UPDATE student SET first_name = %s, middle_name = %s, last_name = %s, sex = %s, course = %s, status = %s WHERE id_no = %s"
        cursor.execute(update_query, (
        updated_first_name, updated_middle_name, updated_last_name, updated_sex, updated_course, updated_status, updated_id_no))
        conn.commit()
        conn.close()

        # Close the update window
        update_window.destroy()

        # Refresh the table with updated data
        populate_treeview(table)

    def populate_treeview(table):
        # Delete existing items in the Treeview
        for row in table.get_children():
            table.delete(row)

        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM student")
        rows = cursor.fetchall()
        conn.commit()
        conn.close()

        for row in rows:
            table.insert("", "end", values=row)
    def search_data():
        search_query = search_entry.get().strip().lower()

        for row in table.get_children():
            table.delete(row)

        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM student")
        rows = cursor.fetchall()

        for row in rows:
            if any(search_query in str(cell).strip().lower() for cell in row):
                table.insert("", "end", values=row)

        conn.close()

    def connection():
        conn = pymysql.connect(host="localhost", user="root", password="", database="libtraq_db")
        return conn

    def go_back():
        list_of_students_window.destroy()

    list_of_students_window = tk.Toplevel()
    list_of_students_window.title("List Of Students")
    list_of_students_window.geometry("1366x768")
    list_of_students_window.resizable(height=False, width=False)

    app_title_label = tk.Label(list_of_students_window, text="Library Tracker and Monitoring System Using QR Code", font=("Arial Rounded MT Bold", 30, "bold"))
    app_title_label.place(x=0, y=0, relwidth=0.99, height=100)

    title_label = tk.Label(list_of_students_window, text="List of Student", bg="gray", font=("Arial", 24, "bold"), borderwidth=1, relief="solid")
    title_label.place(x=0, y=102, relwidth=0.99, height=80)

    back_icon = ImageTk.PhotoImage(Image.open("images/back_icon.png"))
    back_button = tk.Button(list_of_students_window, image=back_icon, command=go_back, bd=0)
    back_button.place(x=10, y=123)

    search_image = Image.open("images/search_icon.png")
    search_photo = ImageTk.PhotoImage(search_image)

    search_button = tk.Button(list_of_students_window, image=search_photo, command=search_data)
    search_button.place(x=1270, y=190)

    search_entry = tk.Entry(list_of_students_window, font=("Arial", 25), width=69, bg="Light Grey")
    search_entry.place(x=10, y=205)

    table_frame = ttk.Frame(list_of_students_window)
    table_frame.place(x=5, y=270)

    table_columns = ("ID Number", "First Name", "Middle Name", "Last Name", "Sex", "Course", "Status", "QR Code", "Photo")
    table = ttk.Treeview(table_frame, columns=table_columns, show="headings", height=20)
    table.place(x=1000, y=1000)

    scrollbar = ttk.Scrollbar(list_of_students_window, orient="vertical", command=table.yview)
    table.configure(yscrollcommand=scrollbar.set)
    table.pack(fill="both", expand=False)
    scrollbar.pack(side="right", fill="y")

    table.bind("<Double-1>", update_student)

    for col in table_columns:
        table.heading(col, text=col)
        table.column(col, width=100)

    # Define column headers
    table.heading("ID Number", text="ID Number")
    table.heading("First Name", text="First Name")
    table.heading("Middle Name", text="Middle Name")
    table.heading("Last Name", text="Last Name")
    table.heading("Sex", text="Sex")
    table.heading("Course", text="Course")
    table.heading("Status", text="Status")
    table.heading("QR Code", text="QR Code")
    table.heading("Photo", text="Photo")

    # Set column widths
    table.column("ID Number", width=148, anchor="center")
    table.column("First Name", width=149, anchor="center")
    table.column("Middle Name", width=149, anchor="center")
    table.column("Last Name", width=149, anchor="center")
    table.column("Sex", width=148, anchor="center")
    table.column("Course", width=148, anchor="center")
    table.column("Status", width=148, anchor="center")
    table.column("QR Code", width=149, anchor="center")
    table.column("Photo", width=149, anchor="center")

    # Populate the Treeview with data
    populate_treeview(table)

    list_of_students_window.mainloop()

def open_generate_report_window():
    generate_report_window = tk.Toplevel()
    generate_report_window.title("Generate Report")
    generate_report_window.geometry("1366x768")
    generate_report_window.resizable(height=False, width=False)

    def refreshTable():
        for date in my_tree.get_children():
            my_tree.delete(date)

        for array in read():
            my_tree.insert(parent='', index='end', iid=array, values=array, tag='orow')
            my_tree.place(x=10, y=300)

    def read():
        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM library_attendance")
        results = cursor.fetchall()
        conn.close()
        return results

    def search_data():
        search_query = search_entry.get().strip().lower()

        for row in my_tree.get_children():
            my_tree.delete(row)

        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM library_attendance")
        rows = cursor.fetchall()

        for row in rows:
            if any(search_query in str(cell).strip().lower() for cell in row):
                my_tree.insert("", "end", values=row)

        conn.close()

    def connection():
        conn = pymysql.connect(host="localhost", user="root", password="", database="libtraq_db")
        return conn

    my_tree = ttk.Treeview(generate_report_window, height=20)

    my_tree['columns'] = ("ID Number", "First Name", "Middle Name", "Last Name", "Course", "Purpose", "Date & Time", "Semester")

    my_tree.column("#0", width=0, stretch=tk.NO)
    my_tree.column("ID Number", anchor="center", width=150)
    my_tree.column("First Name", anchor="center", width=150)
    my_tree.column("Middle Name", anchor="center", width=150)
    my_tree.column("Last Name", anchor="center", width=150)
    my_tree.column("Course", anchor="center", width=140)
    my_tree.column("Purpose", anchor="center", width=190)
    my_tree.column("Date & Time", anchor="center", width=190)
    my_tree.column("Semester", anchor="center", width=190)

    my_tree.heading("ID Number", text="ID Number", anchor="center")
    my_tree.heading("First Name", text="First Name", anchor="center")
    my_tree.heading("Middle Name", text="Middle Name", anchor="center")
    my_tree.heading("Last Name", text="Last Name", anchor="center")
    my_tree.heading("Course", text="Course", anchor="center")
    my_tree.heading("Purpose", text="Purpose", anchor="center")
    my_tree.heading("Date & Time", text="Date & Time", anchor="center")
    my_tree.heading("Semester", text="Semester", anchor="center")

    my_tree.pack()

    refreshTable()
    def go_back():
        generate_report_window.destroy()
    def generate_report():
        selected_courses = course_var.get()

    def print_report():
        selected_courses = course_var.get()

        # Create a Word document
        filename = "report.docx"
        document = Document()

        # Add content to the Word document (customize this part as needed)
        document.add_heading("Library Attendance Report", 0)
        document.add_paragraph("Selected Semester: " + semester_var.get())
        document.add_paragraph("Selected Courses: " + course_var.get())
        document.add_paragraph("Generated by: Mariter Asur")

        # Create a table with headers
        table = document.add_table(rows=1, cols=7)
        table.allow_autofit = False  # Prevent table from auto-adjusting column widths

        table.style = 'Table Grid'  # Apply a table style

        # Add table headers
        table_headers = table.rows[0].cells
        table_headers[0].text = "ID Number"
        table_headers[1].text = "First Name"
        table_headers[2].text = "Middle Name"
        table_headers[3].text = "Last Name"
        table_headers[4].text = "Course"
        table_headers[5].text = "Purpose"
        table_headers[6].text = "Date & Time"

        # Get data from the Treeview and add it to the Word document table
        for item in my_tree.get_children():
            values = my_tree.item(item, 'values')
            data = list(map(str, values))
            table.add_row().cells[0].text = data[0]
            table.rows[-1].cells[1].text = data[1]
            table.rows[-1].cells[2].text = data[2]
            table.rows[-1].cells[3].text = data[3]
            table.rows[-1].cells[4].text = data[4]
            table.rows[-1].cells[5].text = data[5]
            table.rows[-1].cells[6].text = data[6]

        # Save the Word document
        document.save(filename)

        # Display a messagebox
        tk.messagebox.showinfo("Print Successful", "The report has been generated and saved successfully.")

        # Move the file to the "Downloads" folder (change the path as needed)
        downloads_folder = os.path.expanduser("~") + "/Downloads"
        new_filepath = os.path.join(downloads_folder, filename)
        shutil.move(filename, new_filepath)

        print(f"Report generated and saved as {new_filepath}")

    def fetch_and_display_records(dummy_arg=None):
        selected_semester = semester_var.get()
        selected_course = course_var.get()

        # Clear the Treeview widget
        for row in my_tree.get_children():
            my_tree.delete(row)

        # Fetch records from the database based on the selected semester and course
        conn = connection()
        cursor = conn.cursor()

        # Define SQL query based on selected semester and course
        semester_query = ""
        if selected_semester == "1st Semester":
            # Fetch records for August to December in the year 2023
            semester_query = "DATE_FORMAT(date_and_time, '%Y-%m') >= '2023-08' AND DATE_FORMAT(date_and_time, '%Y-%m') <= '2023-12' AND academic_year = 2023"
        elif selected_semester == "2nd Semester":
            # Fetch records for January to June in the year 2024
            semester_query = "DATE_FORMAT(date_and_time, '%Y-%m') >= '2024-01' AND DATE_FORMAT(date_and_time, '%Y-%m') <= '2024-06' AND academic_year = 2024"
        elif selected_semester == "Whole School Year":
            # Fetch records for both semesters
            semester_query = "(DATE_FORMAT(date_and_time, '%Y-%m') >= '2023-08' AND DATE_FORMAT(date_and_time, '%Y-%m') <= '2023-12' AND academic_year = 2023) OR (DATE_FORMAT(date_and_time, '%Y-%m') >= '2024-01' AND DATE_FORMAT(date_and_time, '%Y-%m') <= '2024-06' AND academic_year = 2024)"

        if selected_course == "All Courses":
            course_query = ""
        else:
            course_query = f"course = '{selected_course}'"

        # Construct the full SQL query based on both semester and course
        if semester_query and course_query:
            query = f"SELECT * FROM library_attendance WHERE {semester_query} AND {course_query}"
        elif semester_query:
            query = f"SELECT * FROM library_attendance WHERE {semester_query}"
        elif course_query:
            query = f"SELECT * FROM library_attendance WHERE {course_query}"
        else:
            query = "SELECT * FROM library_attendance"

        cursor.execute(query)
        records = cursor.fetchall()
        conn.close()

        # Display fetched records in the Treeview widget
        for record in records:
            my_tree.insert("", "end", values=record)

    def generate_bar_graph():
        course_counts = {}

        # Count the attendance for each course
        for item in my_tree.get_children():
            values = my_tree.item(item, 'values')
            course = values[4]  # Assuming course is in the 5th column
            if course in course_counts:
                course_counts[course] += 1
            else:
                course_counts[course] = 1

        courses = list(course_counts.keys())
        counts = list(course_counts.values())

        plt.figure(figsize=(10, 6))
        plt.bar(courses, counts, color='blue')
        plt.xlabel('Course')
        plt.ylabel('Attendance Count')
        plt.title('Attendance by Course')
        plt.xticks(rotation=45)
        plt.tight_layout()
        plt.show()

    def create_semester_buttons(semester_var):
        def select_semester(semester):
            semester_var.set(semester)
            fetch_and_display_records()

        semester_label = tk.Label(generate_report_window, text="Select Semester:", font=("Arial", 18))
        semester_label.place(x=0, y=200)

        semester_buttons = []
        for semester in ["1st Semester 2023-2024", "2nd Semester 2023-2024", "Whole School Year"]:
            button = tk.Button(generate_report_window, text=semester, font=("Arial", 16),
                               command=lambda sem=semester: select_semester(sem))
            semester_buttons.append(button)

        for i, button in enumerate(semester_buttons):
            button.place(x=180 + i * 200, y=200)

    def create_course_buttons(course_var):
        def select_course(course):
            course_var.set(course)
            fetch_and_display_records()

        course_label = tk.Label(generate_report_window, text="Select Courses:", font=("Arial", 18))
        course_label.place(x=0, y=240)

        course_buttons = []
        for course in ["BSA", "BSED", "BEED", "BSHM", "BSOA", "BSIT", "All Courses"]:
            button = tk.Button(generate_report_window, text=course, font=("Arial", 16),
                               command=lambda crs=course: select_course(crs))
            course_buttons.append(button)

        for i, button in enumerate(course_buttons):
            button.place(x=180 + i * 120, y=240)

    # Create buttons for selecting semester and courses with the respective variables
    semester_var = tk.StringVar(generate_report_window)
    create_semester_buttons(semester_var)

    course_var = tk.StringVar(generate_report_window)
    create_course_buttons(course_var)

    app_title_label = tk.Label(generate_report_window, text="Library Tracker and Monitoring System Using QR Code", font=("Arial Rounded MT Bold", 30, "bold"))
    app_title_label.place(x=0, y=0, relwidth=0.99, height=100)

    title_label = tk.Label(generate_report_window, text="Generate Report", bg="gray", font=("Arial", 24, "bold"), borderwidth=1, relief="solid")
    title_label.place(x=0, y=102, relwidth=0.99, height=80)

    title_label = tk.Label(generate_report_window, text="Generate Report", bg="gray", font=("Arial", 24, "bold"), borderwidth=1, relief="solid")
    title_label.place(x=0, y=102, relwidth=0.99, height=80)


    graph_button = tk.Button(generate_report_window, bg="orange", text="Graph", font=("Arial", 16),command=generate_bar_graph)
    graph_button.place(x=935, y=200)

    print_button = tk.Button(generate_report_window, text="Print", bg="red" ,font=("Arial", 16),command=print_report)
    print_button.place(x=1170, y=200)

    back_icon = ImageTk.PhotoImage(Image.open("images/back_icon.png"))
    back_button = tk.Button(generate_report_window, image=back_icon, command=go_back, bd=0)
    back_button.place(x=10, y=123)

    search_image = Image.open("images/search_icon.png")
    search_photo = ImageTk.PhotoImage(search_image)

    search_button = tk.Button(generate_report_window, image=search_photo, command=search_data)
    search_button.place(x=1245, y=221)

    #search_entry = tk.Entry(generate_report_window, font=("Arial", 25), width=68, bg="Light Grey")
    #search_entry.place(x=10, y=250)

    generate_report_window.mainloop()

def open_library_utilization_window():
    library_utilization_window = tk.Toplevel()
    library_utilization_window.title("Library Utilization")
    library_utilization_window.geometry("1366x768")
    library_utilization_window.resizable(height=False, width=False)

    def go_back():
        library_utilization_window.destroy()

    def fetch_data_and_create_graph():
        try:
            connect = pymysql.connect(host='localhost', user='root', password="", database='libtraq_db')
            cursor = connect.cursor()

            cursor.execute("SELECT course, COUNT(*) FROM library_attendance GROUP BY course")
            results = cursor.fetchall()
            connect.close()

            if not results:
                messagebox.showerror("Error", "No attendance data found.")
                return

            courses, counts = zip(*results)
            total_count = sum(counts)

            percentages = [count / total_count * 100 for count in counts]

            fig, ax = plt.subplots(figsize=(6, 4))
            bars = ax.bar(courses, counts, color='blue')

            ax.set_xlabel('Course')
            ax.set_ylabel('Attendance Count')
            ax.set_title('Attendance by Course')
            ax.tick_params(axis='x', rotation=45)

            # Add percentages on top of the bars
            for bar, percentage in zip(bars, percentages):
                height = bar.get_height()
                ax.annotate(f'{percentage:.2f}%', (bar.get_x() + bar.get_width() / 2, height),
                            ha='center', va='bottom')

            canvas = FigureCanvasTkAgg(fig, master=library_utilization_window)
            canvas.get_tk_widget().place(x=20, y=320)
            canvas.draw()
        except pymysql.Error as e:
            messagebox.showerror("Database Error", f"Error: {e}")


    def generate_leaderboard():
        try:
            connect = pymysql.connect(host='localhost', user='root', password="", database='libtraq_db')
            cursor = connect.cursor()

            cursor.execute(
                "SELECT first_name, middle_name, last_name, course, purpose, COUNT(*) as attendance_count FROM library_attendance GROUP BY first_name, middle_name, last_name, course, purpose ORDER BY attendance_count DESC")
            results = cursor.fetchall()
            connect.close()

            if not results:
                messagebox.showerror("Error", "No attendance data found.")
                return

            leaderboard_label = tk.Label(library_utilization_window, text="Library Leaderboard",
                                         font=("Arial", 18, "bold"))
            leaderboard_label.place(x=800, y=200)

            # Create a Treeview widget for the table
            leaderboard_tree = ttk.Treeview(library_utilization_window,
                                            columns=("Name", "Course", "Purpose", "Attendance Count"),
                                            show="headings", height=22)
            leaderboard_tree.heading("Name", text="Name", anchor="center")  # Center-align the column heading
            leaderboard_tree.heading("Course", text="Course", anchor="center")
            leaderboard_tree.heading("Purpose", text="Purpose", anchor="center")
            leaderboard_tree.heading("Attendance Count", text="Attendance Count", anchor="center")

            # Style for the column headings (make them larger and bold)
            style = ttk.Style()
            style.configure("Treeview.Heading", font=("Arial", 15, "bold"))

            # Set the column widths
            leaderboard_tree.column("Name", width=170)  # Adjust the width as needed
            leaderboard_tree.column("Course", width=150)  # Adjust the width as needed
            leaderboard_tree.column("Purpose", width=160)  # Adjust the width as needed
            leaderboard_tree.column("Attendance Count", width=180)  # Adjust the width as needed

            leaderboard_tree.place(x=650, y=250)

            # Populate the table with all attendees
            for index, (first_name, middle_name, last_name, course, purpose, attendance_count) in enumerate(results,
                                                                                                            start=1):
                name = f"{first_name} {middle_name} {last_name}"
                leaderboard_tree.insert("", "end", values=(name, course, purpose, attendance_count))

        except pymysql.Error as e:
            messagebox.showerror("Database Error", f"Error: {e}")

    fetch_data_and_create_graph()
    generate_leaderboard()

    app_title_label = tk.Label(library_utilization_window, text="Library Tracker and Monitoring System Using QR Code",
                               font=("Arial Rounded MT Bold", 30, "bold"))
    app_title_label.place(x=0, y=0, relwidth=0.99, height=100)

    title_label = tk.Label(library_utilization_window, text="Library Utilization", bg="gray",
                           font=("Rockwell", 24, "bold"), borderwidth=1, relief="solid")
    title_label.place(x=0, y=102, relwidth=0.99, height=80)

    back_icon = ImageTk.PhotoImage(Image.open("images/back_icon.png"))
    back_button = tk.Button(library_utilization_window, image=back_icon, command=go_back, bd=0)
    back_button.place(x=10, y=123)

    refresh_button = tk.Button(library_utilization_window, bg="orange", text="Leaderboard", font=("Arial", 16),
                               command=lambda: [fetch_data_and_create_graph(), generate_leaderboard()])
    refresh_button.place(x=25, y=200)

    library_utilization_window.mainloop()
def open_about_us_window():
    about_us_window = tk.Toplevel()
    about_us_window.title("About Us")
    about_us_window.geometry("1366x768")
    about_us_window.resizable(height=False, width=False)

    def go_back():
        about_us_window.destroy()

    background_image = Image.open("images/about_us_background.png")
    background_image = background_image.resize((1366, 768), Image.LANCZOS)
    background_photo = ImageTk.PhotoImage(background_image)

    background_label = tk.Label(about_us_window, image=background_photo)
    background_label.place(x=0, y=0, relwidth=1, relheight=1)
    about_us_window.background = background_photo

    back_button = tk.Button(about_us_window, text="Back", command=go_back)
    back_button.place(relx=0, rely=0, anchor=tk.NW)

    about_us_window.mainloop()

def on_entry_click(event):
    if pin_entry.get() == hint_text:
        pin_entry.delete(0, "end")
        pin_entry.config(fg='black')

def on_focus_out(event):
    if not pin_entry.get():
        pin_entry.insert(0, hint_text)
        pin_entry.config(fg='grey')


if __name__ == '__main__':
    window = tk.Tk()
    window.title("LibTraQ: Library Tracker and Monitoring System using QR Code")
    window.geometry("1366x768")
    window.resizable(height=False, width=False)
    hint_text = "Enter PIN"

    background_image = tk.PhotoImage(file="images/admin_background.png")

    background_label = tk.Label(window, image=background_image)
    background_label.place(x=0, y=0, relwidth=1, relheight=1)

    pin_label = tk.Label(window, text="Enter PIN:", font=("Arial", 26), justify='center')
    pin_label.pack(pady=(540, 10))

    pin_entry = tk.Entry(window, show="*", font=("Arial", 30), fg='grey', justify='center')
    pin_entry.insert(0, hint_text)
    pin_entry.bind("<FocusIn>", on_entry_click)
    pin_entry.bind("<FocusOut>", on_focus_out)
    pin_entry.pack(pady=10)
    pin_entry.bind("<Return>", verify_pin)
    pin_entry.focus()

    login_button = tk.Button(window, text="Login", font=("Arial", 24), command=verify_pin)
    login_button.pack(pady=0)

    window.mainloop()
